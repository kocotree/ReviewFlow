"""Build one deterministic review PDF with source separator pages."""

from __future__ import annotations

import logging
import os
import re
import shutil
import subprocess
import unicodedata
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from pypdf import PdfReader, PdfWriter

try:
    # pdfplumber (already a project dependency) installs Pillow.  Keep this
    # optional so a minimal runtime still has the DOCX text fallback.
    from PIL import Image, ImageDraw, ImageFont
except ImportError:  # pragma: no cover - exercised only in a stripped runtime
    Image = ImageDraw = ImageFont = None

from app.docx_convert import attachment_to_pdf
from app.models.content import MaterialGroup, PdfMaterial

logger = logging.getLogger(__name__)

MAX_SEPARATOR_TITLE_CHARS = 120
_GROUP_ORDER = {
    MaterialGroup.ONLINE_DOC: 0,
    MaterialGroup.ATTACHMENT: 1,
}

_URL_RE = re.compile(r"(?:https?|lark|feishu)://\S+", re.IGNORECASE)
_TECHNICAL_METADATA_RE = re.compile(
    r"\b(?:tmp[_-]?url|(?:file|tenant|access|app|record|table|doc|obj|wiki|node)"
    r"[_-]?token)\s*[:=]\s*\S+",
    re.IGNORECASE,
)

_FONT_CANDIDATES: tuple[tuple[str, tuple[str, ...]], ...] = (
    (
        "PingFang SC",
        (
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/PingFangSC-Regular.otf",
        ),
    ),
    ("Hiragino Sans GB", ("/System/Library/Fonts/Hiragino Sans GB.ttc",)),
    ("STHeiti", ("/System/Library/Fonts/STHeiti Medium.ttc",)),
    (
        "Noto Sans CJK SC",
        (
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        ),
    ),
    (
        "Microsoft YaHei",
        (str(Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts/msyh.ttc"),),
    ),
    (
        "SimHei",
        (str(Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts/simhei.ttf"),),
    ),
)


class PdfBundleError(RuntimeError):
    """Raised when a complete, trustworthy review bundle cannot be produced."""


def sanitize_source_title(
    title: str,
    *,
    max_chars: int = MAX_SEPARATOR_TITLE_CHARS,
) -> str:
    """Clean untrusted source labels before placing them in the PDF.

    Control characters, URLs, and common token-like technical metadata are
    omitted.  The visible title is then whitespace-normalized and truncated.
    """
    if max_chars < 4:
        raise ValueError("max_chars must be at least 4")

    value = unicodedata.normalize("NFKC", str(title or ""))
    value = _URL_RE.sub(" ", value)
    value = _TECHNICAL_METADATA_RE.sub(" ", value)
    value = "".join(
        " " if char.isspace() else char
        for char in value
        if not unicodedata.category(char).startswith("C") or char.isspace()
    )
    value = re.sub(r"\s+", " ", value).strip(" ._-—")
    if not value:
        value = "未命名材料"
    if len(value) > max_chars:
        value = value[: max_chars - 3].rstrip() + "..."
    return value


def _ordered_materials(materials: Iterable[PdfMaterial]) -> list[PdfMaterial]:
    """Sort by group/sequence while preserving input order for equal keys."""
    indexed = list(enumerate(materials))
    return [
        material
        for _, material in sorted(
            indexed,
            key=lambda pair: (
                _GROUP_ORDER[pair[1].group],
                pair[1].sequence,
                pair[0],
            ),
        )
    ]


async def build_pdf_bundle(materials: Iterable[PdfMaterial]) -> bytes:
    """Merge materials into one PDF, each preceded by a source separator page.

    Ordering is explicit and deterministic: online documents precede
    attachments, each group is ordered by ``sequence``, and duplicate sequence
    values retain their input order.
    """
    ordered = _ordered_materials(materials)
    if not ordered:
        raise PdfBundleError("无法生成总 PDF：没有文件型材料")

    writer = PdfWriter()
    # PdfReader can retain references to its byte stream until writer.write().
    # Keep all streams/readers alive for the complete merge operation.
    retained: list[tuple[BytesIO, PdfReader]] = []

    try:
        for material in ordered:
            safe_title = sanitize_source_title(material.title)
            separator_pdf = await _create_source_separator_pdf(
                material.group.display_name,
                safe_title,
            )
            _append_pdf(
                writer,
                separator_pdf,
                retained,
                label=f"{safe_title} 的来源分隔页",
            )
            _append_pdf(
                writer,
                material.pdf_bytes,
                retained,
                label=safe_title,
            )

        output = BytesIO()
        writer.write(output)
        bundle = output.getvalue()
        if not bundle:
            raise PdfBundleError("总 PDF 合并失败：生成结果为空")
        return bundle
    except PdfBundleError:
        raise
    except Exception as exc:
        raise PdfBundleError(f"总 PDF 合并失败：{exc}") from exc
    finally:
        close = getattr(writer, "close", None)
        if callable(close):
            close()


def _append_pdf(
    writer: PdfWriter,
    pdf_bytes: bytes,
    retained: list[tuple[BytesIO, PdfReader]],
    *,
    label: str,
) -> None:
    """Validate and append all pages from one PDF input."""
    if not isinstance(pdf_bytes, bytes) or not pdf_bytes:
        raise PdfBundleError(f"材料 PDF 无效或为空：{label}")

    stream = BytesIO(pdf_bytes)
    try:
        reader = PdfReader(stream, strict=False)
        if reader.is_encrypted:
            raise PdfBundleError(f"材料 PDF 已加密，无法合并：{label}")
        page_count = len(reader.pages)
    except PdfBundleError:
        raise
    except Exception as exc:
        raise PdfBundleError(f"材料 PDF 损坏或无法解析：{label}") from exc

    if page_count == 0:
        raise PdfBundleError(f"材料 PDF 不包含有效页面：{label}")

    try:
        for page in reader.pages:
            writer.add_page(page)
    except Exception as exc:
        raise PdfBundleError(f"材料 PDF 页面合并失败：{label}") from exc
    retained.append((stream, reader))


async def _create_source_separator_pdf(source_type: str, title: str) -> bytes:
    """Create a one-page DOCX separator and convert it with LibreOffice."""
    docx_bytes = _create_source_separator_docx(source_type, title)
    pdf = await attachment_to_pdf(docx_bytes, "source-separator.docx")
    if not pdf:
        raise PdfBundleError(f"来源分隔页生成失败：{title}")
    return pdf


def _create_source_separator_docx(source_type: str, title: str) -> bytes:
    """Create a restrained A4 separator page with a CJK-capable font."""
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(24)
    section.right_margin = Mm(24)

    separator_image = _render_source_separator_image(source_type, title)
    if separator_image:
        # LibreOffice builds can differ in how they map platform CJK fonts into
        # exported PDFs.  Rasterizing only this small separator heading with an
        # installed system font makes Chinese labels reliable while DOCX still
        # controls the A4 page and LibreOffice remains the PDF renderer.
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(165)
        run = paragraph.add_run()
        run.add_picture(BytesIO(separator_image), width=Mm(162))
        document.inline_shapes[-1]._inline.docPr.set(
            "descr", f"{source_type}: {title}"
        )
    else:
        _add_text_separator_fallback(document, source_type, title)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _render_source_separator_image(source_type: str, title: str) -> bytes | None:
    """Render separator text with a concrete system CJK font when possible."""
    font_path = _preferred_cjk_font_path()
    if Image is None or ImageDraw is None or ImageFont is None or font_path is None:
        return None

    try:
        canvas = Image.new("RGB", (1400, 500), "white")
        draw = ImageDraw.Draw(canvas)
        label_font = ImageFont.truetype(str(font_path), 42)
        draw.text(
            (700, 58),
            source_type,
            font=label_font,
            fill=(91, 103, 120),
            anchor="ma",
        )

        title_font, title_lines = _fit_title(draw, title, font_path)
        line_box = draw.textbbox((0, 0), "国Ag", font=title_font)
        line_height = line_box[3] - line_box[1] + 20
        title_height = line_height * len(title_lines)
        y = 165 + max(0, (285 - title_height) // 2)
        for line in title_lines:
            draw.text(
                (700, y),
                line,
                font=title_font,
                fill=(31, 41, 55),
                anchor="ma",
                stroke_width=1,
                stroke_fill=(31, 41, 55),
            )
            y += line_height

        output = BytesIO()
        canvas.save(output, format="PNG", optimize=True)
        return output.getvalue()
    except Exception:
        logger.warning("系统中文字体渲染失败，回退 DOCX 文本", exc_info=True)
        return None


def _fit_title(draw, title: str, font_path: Path):
    """Select a title size and character-wrap it inside the image bounds."""
    for size in range(72, 31, -4):
        font = ImageFont.truetype(str(font_path), size)
        lines = _wrap_text(draw, title, font, max_width=1240)
        line_box = draw.textbbox((0, 0), "国Ag", font=font)
        line_height = line_box[3] - line_box[1] + 20
        if len(lines) <= 5 and line_height * len(lines) <= 285:
            return font, lines

    font = ImageFont.truetype(str(font_path), 32)
    return font, _wrap_text(draw, title, font, max_width=1240)[:5]


def _wrap_text(draw, text: str, font, *, max_width: int) -> list[str]:
    """Wrap mixed CJK/Latin titles without relying on language tokenization."""
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if current and width > max_width:
            lines.append(current.rstrip())
            current = char.lstrip()
        else:
            current = candidate
    if current or not lines:
        lines.append(current.rstrip())
    return lines


def _add_text_separator_fallback(
    document: Document,
    source_type: str,
    title: str,
) -> None:
    """Build a normal text separator when Pillow/system font files are absent."""
    font_name = _preferred_cjk_font()
    normal_style = document.styles["Normal"]
    normal_style.font.name = font_name
    normal_rpr = normal_style._element.get_or_add_rPr()
    normal_rpr.get_or_add_rFonts().set(qn("w:eastAsia"), font_name)

    source_paragraph = document.add_paragraph()
    source_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_paragraph.paragraph_format.space_before = Pt(190)
    source_run = source_paragraph.add_run(source_type)
    _set_run_font(source_run, font_name, size=Pt(14))
    source_run.font.color.rgb = RGBColor(91, 103, 120)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(16)
    title_paragraph.paragraph_format.space_after = Pt(0)
    title_run = title_paragraph.add_run(title)
    _set_run_font(title_run, font_name, size=Pt(24), bold=True)
    title_run.font.color.rgb = RGBColor(31, 41, 55)


def _set_run_font(run, font_name: str, *, size: Pt, bold: bool = False) -> None:
    """Set both Latin and East Asian font names for LibreOffice rendering."""
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), font_name
    )


@lru_cache(maxsize=1)
def _preferred_cjk_font() -> str:
    """Choose an installed CJK font, falling back to the system zh font."""
    for family, paths in _FONT_CANDIDATES:
        if any(Path(path).is_file() for path in paths):
            return family

    fc_match = shutil.which("fc-match")
    if fc_match:
        try:
            result = subprocess.run(
                [fc_match, "-f", "%{family}", "sans-serif:lang=zh"],
                check=True,
                capture_output=True,
                text=True,
                timeout=3,
            )
            family = result.stdout.split(",", 1)[0].strip()
            if family:
                return family
        except (OSError, subprocess.SubprocessError):
            logger.debug("无法通过 fontconfig 定位中文字体", exc_info=True)

    # LibreOffice will apply its platform fallback if this family is absent.
    return "Arial Unicode MS"


@lru_cache(maxsize=1)
def _preferred_cjk_font_path() -> Path | None:
    """Return a concrete installed CJK font file suitable for Pillow."""
    for _, paths in _FONT_CANDIDATES:
        for path in paths:
            candidate = Path(path)
            if candidate.is_file():
                return candidate
    return None
